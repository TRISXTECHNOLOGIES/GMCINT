from docx import Document
from docx.shared import Inches,Pt
from docx2pdf import convert
from PIL import Image
import os
from datetime import date
#import pythoncom

today_date = date.today().strftime("%d day of %B, %Y")

def consent_form_minor(
        user_id,
        flight_id,
        guardian_name, 
        relation_name, 
        flight_number, 
        airline_name):
    
   #pythoncom.CoInitialize()
    try:
        signature_path = f'uploads/{user_id}/signature_{flight_id}.png'
        output_docx = f'uploads/{user_id}/consent_form_minor_{flight_id}.docx'
        # Load the document
        doc = Document('consent_form_minor.docx')
        
        # Replace placeholders with dynamic content
        for para in doc.paragraphs:
            if "[guardian_name]" in para.text:
                para.text = para.text.replace("[guardian_name]", guardian_name)
            if "[relation_name]" in para.text:
                para.text = para.text.replace("[relation_name]", relation_name)
            if "[flight_number]" in para.text:
                para.text = para.text.replace("[flight_number]", flight_number)
            if "[airline_name]" in para.text:
                para.text = para.text.replace("[airline_name]", airline_name)
            if "[date]" in para.text:
                para.text = para.text.replace("[date]", today_date)
            if "[signature]" in para.text:
                # Insert the signature
                try:
                    para.clear()
                    run = para.add_run()
                    run.add_picture(signature_path, width=Inches(2), height=Inches(1))
                except Exception as e:
                    print(f"Error inserting signature: {e}")
            if "[title]" in para.text:
                para.text = para.text.replace("[title]", 'GET MY CLAIM')
            if "[guardian_name_footer]" in para.text:
                para.clear()
                pts = len(guardian_name)
                pts = int((144/pts)*2)
                if pts>=12:
                    pts=12
                # Add new run for airline name
                run = para.add_run(guardian_name)
                run.font.size = Pt(pts)  # Set font size to 16 points (adjust as needed)
                run.bold = False

        # Save the updated document as .docx
        doc.save(output_docx)
        print(f"Document saved as {output_docx}")

        # Convert the .docx to PDF
        pdf_output = output_docx.replace(".docx", ".pdf")
        convert(output_docx, pdf_output)  # Convert .docx to .pdf
        os.remove(output_docx)
        print(f"Document converted to PDF: {pdf_output}")
    except Exception as e:
        print(f"An error occurred: {e}")
    

def consent_form_major(
        user_id,
        flight_id,
        guardian_name, 
        relation_name,
        flight_number, 
        airline_name):
    
   #pythoncom.CoInitialize()
    try:
        signature_path = f'uploads/{user_id}/signature_{flight_id}.png'
        output_docx = f'uploads/{user_id}/consent_form_major_{flight_id}.docx'
        # Load the document
        doc = Document('consent_form_major.docx')
        
        # Replace placeholders with dynamic content
        for para in doc.paragraphs:
            if "[guardian_name]" in para.text:
                para.text = para.text.replace("[guardian_name]", guardian_name)
            if "[relation_name]" in para.text:
                para.text = para.text.replace("[relation_name]", relation_name)
            if "[flight_number]" in para.text:
                para.text = para.text.replace("[flight_number]", flight_number)
            if "[airline_name]" in para.text:
                para.text = para.text.replace("[airline_name]", airline_name)
            if "[date]" in para.text:
                para.text = para.text.replace("[date]", today_date)
            if "[signature]" in para.text:
                # Insert the signature
                try:
                    para.clear()
                    run = para.add_run()
                    run.add_picture(signature_path, width=Inches(2), height=Inches(1))
                except Exception as e:
                    print(f"Error inserting signature: {e}")
            if "[title]" in para.text:
                para.text = para.text.replace("[title]", 'GET MY CLAIM')
            if "[guardian_name_footer]" in para.text:
                para.clear()
                pts = len(guardian_name)
                pts = int((144/pts)*2)
                if pts>=12:
                    pts=12
                # Add new run for airline name
                run = para.add_run(guardian_name)
                run.font.size = Pt(pts)  # Set font size to 16 points (adjust as needed)
                run.bold = False

        # Save the updated document as .docx
        doc.save(output_docx)
        print(f"Document saved as {output_docx}")

        # Convert the .docx to PDF
        pdf_output = output_docx.replace(".docx", ".pdf")
        convert(output_docx, pdf_output)  # Convert .docx to .pdf
        os.remove(output_docx)
        print(f"Document converted to PDF: {pdf_output}")
    except Exception as e:
        print(f"An error occurred: {e}")


def consent_form_individual(
        user_id,
        flight_id,
        guardian_name, 
        flight_number, 
        airline_name):
    
   #pythoncom.CoInitialize()
    try:
        signature_path = f'uploads/{user_id}/signature_{flight_id}.png'
        output_docx = f'uploads/{user_id}/consent_form_individual_{flight_id}.docx'
        # Load the document
        doc = Document('consent_form_individual.docx')
        
        # Replace placeholders with dynamic content
        for para in doc.paragraphs:
            if "[guardian_name]" in para.text:
                para.text = para.text.replace("[guardian_name]", guardian_name)
            if "[flight_number]" in para.text:
                para.text = para.text.replace("[flight_number]", flight_number)
            if "[airline_name]" in para.text:
                para.text = para.text.replace("[airline_name]", airline_name)
            if "[date]" in para.text:
                para.text = para.text.replace("[date]", today_date)
            if "[signature]" in para.text:
                # Insert the signature
                try:
                    para.clear()
                    run = para.add_run()
                    run.add_picture(signature_path, width=Inches(2), height=Inches(1))
                except Exception as e:
                    print(f"Error inserting signature: {e}")
            if "[title]" in para.text:
                para.text = para.text.replace("[title]", 'GET MY CLAIM')
            if "[guardian_name_footer]" in para.text:
                para.clear()
                pts = len(guardian_name)
                pts = int((144/pts)*2)
                if pts>=12:
                    pts=12
                # Add new run for airline name
                run = para.add_run(guardian_name)
                run.font.size = Pt(pts)  # Set font size to 16 points (adjust as needed)
                run.bold = False

        # Save the updated document as .docx
        doc.save(output_docx)
        print(f"Document saved as {output_docx}")

        # Convert the .docx to PDF
        pdf_output = output_docx.replace(".docx", ".pdf")
        convert(output_docx, pdf_output)  # Convert .docx to .pdf
        os.remove(output_docx)
        print(f"Document converted to PDF: {pdf_output}")
    except Exception as e:
        print(f"An error occurred: {e}")
    
