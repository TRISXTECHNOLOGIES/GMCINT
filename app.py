import os, re
import base64
from io import BytesIO
from PIL import Image
from flask import Flask, render_template, request, redirect, url_for, session, flash,jsonify, send_from_directory
from flask_sqlalchemy import SQLAlchemy
from werkzeug.utils import secure_filename
from werkzeug.security import generate_password_hash, check_password_hash
import uuid 
import random
import time
import io
import csv
from flask_migrate import Migrate
import chardet
import pandas as pd
import smtplib
from email.message import EmailMessage
from mailer import partner_client_email
from flask import send_from_directory
import datetime
from werkzeug.security import generate_password_hash

class Config:
    SECRET_KEY = os.environ.get('SECRET_KEY') or 'you-will-never-guess'
    SQLALCHEMY_DATABASE_URI = 'mysql://doadmin:AVNS_OfESOuDWrnHECOFfwJ1@db-mysql-blr1-75277-do-user-17878006-0.l.db.ondigitalocean.com:25060/defaultdb'
    SQLALCHEMY_TRACK_MODIFICATIONS = False
    SQLALCHEMY_POOL_RECYCLE = 280
    SQLALCHEMY_POOL_PRE_PING = True
    SQLALCHEMY_POOL_TIMEOUT = 30
    UPLOAD_FOLDER = 'uploads'
    ALLOWED_EXTENSIONS = {'png', 'jpg', 'jpeg', 'gif', 'pdf'}

# Initialize app
app = Flask(__name__)
app.config.from_object(Config)

# Initialize DB
db = SQLAlchemy(app)
migrate = Migrate(app, db)

# DB initialization function
def init_db():
    with app.app_context():
        db.create_all()


def gen_uid(id_for):
    while True:
        random_id = random.randint(10000, 99999)
        # Check if the generated ID is unique
        if not id_for.query.filter_by(id=random_id).first():
            return random_id
        
def gen_passwd():
    sp_char = ['!', '@', '#', '$', '%', '^', '&', '*', '-', '_', '=', '+', '?' ]
    passwd = str(uuid.uuid4()).replace('-','')[:7]
    passwd = list(passwd + random.choice(sp_char))
    random.shuffle(passwd)
    password = ''.join(passwd)
    return password

def filter_airports(user_input):
    with open('Airports_list.txt', 'r') as list:
        list = list.read()
        lines = list.splitlines()
    airport_data = {}
    current_country = None
    # Iterate over each line of text
    for line in lines:
        line = line.strip()  
        if line.endswith(':'):
            current_country = line[:-1]
            airport_data[current_country] = []  
        elif line and current_country:
            airport_data[current_country].append(line.split(". ", 1)[-1])  # Remove numbering like '1.'
    suggestions = []
    user_input = user_input.lower()  
    found = False
    
    for country, airports in airport_data.items():
        if country.lower().startswith(user_input):
            for airport in airports:
                suggestions.append(airport)
            found = True
        else:
            for airport in airports:
                if airport.lower().startswith(user_input):
                    suggestions.append(airport)
                    found = True
                elif user_input in airport.lower() and len(user_input)>=2:
                    # print(f"{country}:")    
                    # print(f"  {airport}\n")
                    suggestions.append(airport)
                    found = True
    if not found:
        no_data =  "No results found!"
        suggestions.append(no_data)
        return suggestions
    else:
        return suggestions
    
def filter_airlines(user_input):
    with open('Airlines_list.txt','r') as list:
        list = list.readlines()
    suggestions = []
    user_input = user_input.lower()  
    found = False
    for item in list:
        if item.lower().startswith(user_input):
            suggestions.append(item)
            found = True
        elif user_input in item.lower():
            suggestions.append(item)
            found = True
    if not found:
        no_data =  "No results found!"
        suggestions.append(no_data)
        return suggestions
    else:
        return suggestions

def read_csv_with_encoding_detection(file_path):
    """
    Robust CSV reader that handles various encodings
    """
    # First, detect encoding
    with open(file_path, 'rb') as f:
        raw_data = f.read()
        encoding_result = chardet.detect(raw_data)
        detected_encoding = encoding_result['encoding']
    
    # List of encodings to try (in order of preference)
    encodings_to_try = [
        detected_encoding,
        'utf-8',
        'latin-1',
        'iso-8859-1',
        'cp1252',  # Windows encoding
        'utf-16',
        'ascii'
    ]
    
    # Remove None and duplicates
    encodings_to_try = list(dict.fromkeys([enc for enc in encodings_to_try if enc]))
    
    for encoding in encodings_to_try:
        try:
            df = pd.read_csv(file_path, encoding=encoding)
            
            # Clean problematic characters
            for col in df.select_dtypes(include=['object']).columns:
                df[col] = df[col].astype(str).str.replace('\xa0', ' ').str.strip()
            
            print(f"Successfully read CSV with encoding: {encoding}")
            return df
            
        except (UnicodeDecodeError, UnicodeError) as e:
            print(f"Failed to read with encoding {encoding}: {e}")
            continue
    
    raise ValueError("Unable to read CSV file with any supported encoding")

def send_email(to_email, subject, body):
    msg = EmailMessage()
    msg.set_content(body)
    msg['Subject'] = subject
    msg['From'] = 'no-reply@getmyclaims.com'
    msg['To'] = to_email

    with smtplib.SMTP('smtp.gmail.com', 587) as server:
        server.starttls()
        server.login('no-reply@getmyclaims.com', 'tzul utwx lyyn bmtu')  # Use Gmail App Password
        server.send_message(msg)

# Define database models
class User(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    first_name = db.Column(db.String(100))
    last_name = db.Column(db.String(100))
    email = db.Column(db.String(100))
    password = db.Column(db.String(200))
    role = db.Column(db.String(50), default='user')
    # Specify which foreign key to use for this relationship
    flights = db.relationship('Flight', backref='user', lazy=True, foreign_keys='Flight.user_id')
    # Add a new relationship for partner flights
    partner_flights = db.relationship('Flight', backref='partner', lazy=True,foreign_keys='Flight.partner_id')
    assigned_admin_id = db.Column(db.Integer, db.ForeignKey('user.id'))

class Flight(db.Model): 
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'))
    leaving_from = db.Column(db.String(100))
    destination = db.Column(db.String(100))
    connecting_flights = db.Column(db.String(100))
    connecting_leaving_from = db.Column(db.String(100))
    connecting_leaving_from2 = db.Column(db.String(100))
    compensation_flight = db.Column(db.String(100))
    date = db.Column(db.String(50))
    disruption = db.Column(db.String(100))
    delay = db.Column(db.String(20))
    reason = db.Column(db.String(100))
    airline = db.Column(db.String(100))
    airline2 = db.Column(db.String(100))
    airline3 = db.Column(db.String(100))
    flight_number = db.Column(db.String(50))
    flight_date = db.Column(db.String(50))
    address_line1 = db.Column(db.String(200))
    address_line2 = db.Column(db.String(200))
    first_name = db.Column(db.String(100))
    last_name = db.Column(db.String(100))
    first_name_passenger_2 = db.Column(db.String(100))
    last_name_passenger_2 = db.Column(db.String(100))
    passenger_2_type = db.Column(db.String(100))
    first_name_passenger_3 = db.Column(db.String(100))
    last_name_passenger_3 = db.Column(db.String(100))
    passenger_3_type = db.Column(db.String(100))
    first_name_passenger_4 = db.Column(db.String(100))
    last_name_passenger_4 = db.Column(db.String(100))
    passenger_4_type = db.Column(db.String(100))
    first_name_passenger_5 = db.Column(db.String(100))
    last_name_passenger_5 = db.Column(db.String(100))
    passenger_5_type = db.Column(db.String(100))
    email = db.Column(db.String(100))
    group_travel = db.Column(db.String(100))
    city = db.Column(db.String(100))
    postal_code = db.Column(db.String(20))
    country = db.Column(db.String(100))
    phone_number = db.Column(db.String(50))
    notification = db.Column(db.String(50))
    booking_number = db.Column(db.String(50))
    
    signature = db.Column(db.LargeBinary)  # Legacy (not currently used but kept for compatibility)
    signature_path = db.Column(db.String(255))  # ✅ New - path to saved signature image
    signed_file_path = db.Column(db.String(255))  # ✅ Final DOCX with signature
    boarding_pass = db.Column(db.String(200))
    contacted_airline = db.Column(db.String(200))
    incident_description = db.Column(db.Text)
    reference_id = db.Column(db.Integer, unique=True)
    password = db.Column(db.String(100))
    status = db.Column(db.String(50), default='pending')  # ✅ Changed default from 'NA' to 'pending'
    partner_id = db.Column(db.Integer, db.ForeignKey('user.id'), nullable=True)
    verification_token = db.Column(db.String(255), nullable=True)  # ✅ Token for claim link
    
    comments = db.relationship('Comment', backref='flight', lazy=True)
    
class Partner(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(100))
    email = db.Column(db.String(100), unique=True)
    password = db.Column(db.String(200))   

# class Flight(db.Model):
#     id = db.Column(db.Integer, primary_key=True)
#     user_id = db.Column(db.Integer, db.ForeignKey('user.id'))
#     leaving_from = db.Column(db.String(100))
#     destination = db.Column(db.String(100))
#     connecting_flights = db.Column(db.String(100))
#     date = db.Column(db.String(50))
#     disruption = db.Column(db.String(100))
#     delay = db.Column(db.String(20))
#     airline = db.Column(db.String(100))
#     flight_number = db.Column(db.String(50))
#     flight_date = db.Column(db.String(50))
#     address_line1 = db.Column(db.String(200))
#     address_line2 = db.Column(db.String(200))
#     first_name = db.Column(db.String(100))
#     first_name_passenger_2 = db.Column(db.String(100))
#     last_name_passenger_2=db.Column(db.String(100))
#     last_name = db.Column(db.String(100))
#     email = db.Column(db.String(100))
#     email_passenger_2=db.Column(db.String(100))
#     group_travel=db.Column(db.String(100))
#     city = db.Column(db.String(100))
#     postal_code = db.Column(db.String(20))
#     state = db.Column(db.String(100))
#     country = db.Column(db.String(100))
#     phone_number = db.Column(db.String(50))
#     booking_number = db.Column(db.String(50))
#     signature = db.Column(db.LargeBinary)
#     boarding_pass = db.Column(db.String(200))
#     contacted_airline = db.Column(db.String(200))
#     incident_description = db.Column(db.Text)
#     status = db.Column(db.String(50),default='NA')
#     comments = db.relationship('Comment', backref='flight', lazy=True)

class Comment(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    user_id = db.Column(db.Integer, db.ForeignKey('user.id'))
    flight_id = db.Column(db.Integer, db.ForeignKey('flight.id'))
    admin_name = db.Column(db.String(50))
    text = db.Column(db.Text)
    timestamp = db.Column(db.DateTime, default=db.func.now())

# Ensure the upload folder exists
if not os.path.exists(app.config['UPLOAD_FOLDER']):
    os.makedirs(app.config['UPLOAD_FOLDER'])

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in app.config['ALLOWED_EXTENSIONS']

@app.template_filter('tojson')
def to_json_filter(obj):
    import json
    return json.dumps(obj)

@app.route('/clear_session')
def clear_session():
    session.clear()
    return "Session cleared!"

@app.route('/', methods=['GET','POST'])
def landing():
    state='login'
    if 'role' in session.keys():
        state='dashboard'
    
    if request.method == 'POST':
        data = dict()
        data['leaving_from'] = ''
        data['destination'] = ''
        data['first_name'] = ''
        data['last_name'] = ''
        data['airline'] = ''
        if 'leaving_from' in request.form:
            leaving_from = request.form.get('leaving_from')
            data['leaving_from'] = leaving_from
        if 'destination' in request.form:
            destination = request.form.get('destination')
            data['destination'] = destination
        if 'boarding_pass' in request.files:
            boarding_pass = request.files.get('boarding_pass')
            if boarding_pass.filename!='':
                if boarding_pass and allowed_file(boarding_pass.filename):
                    path = os.path.join(app.config['UPLOAD_FOLDER'],boarding_pass.filename)
                    boarding_pass.save(path)
                    session['temp_path'] = path
                    print(path)
        if 'temp_path' not in session.keys():
            return render_template('form.html', data=data)
        elif 'temp_path' in session.keys():
            print('temp_path')
            from scan_text import extract_entities
            path = session['temp_path']
            names, airports, airline = extract_entities(path)
            if names:
                if len(names)==2:
                    first_name = names[0]
                    data['first_name'] = first_name
                    last_name = names[1]
                    data['last_name'] = last_name
                else:
                    first_name = names[0]
                    data['first_name'] = first_name
            if airports:
                print(airports)
                if len(airports)==2:
                    leaving_from = airports[0]
                    data['leaving_from'] = leaving_from
                    destination = airports[1]
                    data['destination'] = destination
                else:
                    leaving_from = airports[0]
                    data['leaving_from'] = leaving_from
            if airline:
                data['airline'] = airline
            print(data.items())
            session.pop('temp_path')
        return render_template('form.html', data=data)
        
    return render_template('index.html',state=state)

@app.route('/kyr')
def kyr():
    return render_template('kyr.html')


@app.route('/about')
def about():
    return render_template('about.html')

@app.route('/blog')
def blog():
    return render_template('blog.html')

@app.route('/partner')
def partner():
    return render_template('partner.html')

@app.route('/help')
def help():
    return render_template('help.html')

@app.route('/fees')
def fees():
    return render_template('fees.html')

@app.route('/refer')
def refer():
    return render_template('refer.html')

@app.route('/T&C')
def TC():
    return render_template('T&C.html')

@app.route('/policy')
def policy():
    return render_template('policy.html')

@app.route('/cookies')
def cookies():
    return render_template('cookies.html')

@app.route('/contact')
def contact():
    return render_template('contact.html')

@app.route('/airlines')
def airlines():
    return render_template('airlines.html')

@app.route('/index', methods=['GET', 'POST'])
def index():
    #form data
    if request.method == 'POST':
        print("***********")
        details = request.form.to_dict()
        for key, value in details.items():
            print(key,":",value)
        print("***********\n\n\n")

        first_time_user = False
        if 'user_id' not in session:
            first_name = request.form.get('first_name','John')
            last_name = request.form.get('last_name','Doe')
            email = request.form.get('email').strip().lower()
            if email:
                if User.query.filter_by(email=email).first():
                    flash('Your email is already registerd, login please!')
                    return redirect(url_for('login'))
            
            password = gen_passwd()
            user_id = gen_uid(User)
            
            # user.id = user_id
            session['user_id'] = user_id
            session['username'] = f"{first_name} {last_name}"
            session['role'] = 'user'
            session['email'] = email
            session['password'] = password
            user = User(id=user_id,first_name=first_name, last_name=last_name, email=email, password=generate_password_hash(password),role='user')
            db.session.add(user)
            db.session.commit()
            # db.session.commit()
            first_time_user = True

        user = db.session.get(User, session['user_id']) 
        user_dir_path = os.makedirs(os.path.join(app.config['UPLOAD_FOLDER'], f"{session['user_id']}"), exist_ok=True)
        try:
            if 'leaving_from' in request.form:
                flight = Flight(
                    id = gen_uid(Flight),
                    user_id=user.id,
                    leaving_from=request.form.get('leaving_from'),
                    destination=request.form.get('destination'),
                )
                db.session.add(flight)
                db.session.commit()

                # flight.id = gen_uid(Flight)
                # db.session.commit()
                session['flight_id'] = flight.id
                print("Flight added:", flight)
                print(request.form.get('leaving_from'),request.form.get('destination'))
            flight_id = session.get('flight_id')

            if 'connecting_flights' in request.form and flight_id:
                flight = db.session.get(Flight,flight_id)
                flight.connecting_flights = request.form.get('connecting_flights') 
                db.session.commit()
            if 'connecting_leaving_from' in request.form and flight_id:
                flight = db.session.get(Flight,flight_id)
                flight.connecting_leaving_from = request.form.get('connecting_leaving_from') 
                db.session.commit()
            if 'connecting_leaving_from2' in request.form and flight_id:
                flight = db.session.get(Flight,flight_id)
                flight.connecting_leaving_from2 = request.form.get('connecting_leaving_from2') 
                db.session.commit()
            if 'compensation_flight' in request.form and flight_id:
                flight = db.session.get(Flight,flight_id)
                flight.compensation_flight = request.form.get('compensation_flight') 
                db.session.commit()

            if flight_id:
                flight = db.session.get(Flight,flight_id)
                from datetime import date
                current_date = str(date.today())
                flight.date = current_date
                db.session.commit()
                
                
            if 'disruption' in request.form and flight_id:
                flight = db.session.get(Flight,flight_id)
                flight.disruption = request.form.get('disruption')  # Assuming this is what you meant
                flight.delay = request.form.get('delay')  # Assuming this is what you meant
                flight.reason = request.form.get('reason')
                db.session.commit()
                print("Flight updated with disruption:", flight)

            if 'airline' in request.form and flight_id:
                flight = db.session.get(Flight,flight_id)
                flight.airline = request.form.get('airline')
                flight.airline2 = request.form.get('airline2')
                flight.airline3 = request.form.get('airline3')
                flight.booking_number = request.form.get('booking_number')
                flight.flight_number = request.form.get('flight_number')
                flight.flight_date = request.form.get('flight_date')
                db.session.commit()
                
            
            if 'group_travel' in request.form and flight_id:
                flight = db.session.get(Flight,flight_id)
                flight.group_travel = request.form.get('group_travel')
                db.session.commit()
            minors = []
            majors = []
            if 'first_name_passenger_2' in request.form and flight_id:
                flight = db.session.get(Flight,flight_id)
                flight.first_name_passenger_2 = request.form.get('first_name_passenger_2')
                flight.last_name_passenger_2 = request.form.get('last_name_passenger_2')
                flight.passenger_2_type = request.form.get('passenger_2_type')
                fname = request.form.get('first_name_passenger_2')
                lname = request.form.get('last_name_passenger_2')
                if 'passenger_2_type' in request.form:
                    minors.append(f'{fname} {lname}')
                else:
                    majors.append(f'{fname} {lname}')
                db.session.commit()
                
            if 'first_name_passenger_3' in request.form and flight_id:
                flight = db.session.get(Flight,flight_id)
                flight.first_name_passenger_3 = request.form.get('first_name_passenger_3')
                flight.last_name_passenger_3 = request.form.get('last_name_passenger_3')
                flight.passenger_3_type = request.form.get('passenger_3_type')
                fname = request.form.get('first_name_passenger_3')
                lname = request.form.get('last_name_passenger_3')
                if 'passenger_3_type' in request.form:
                    minors.append(f'{fname} {lname}')
                else:
                    majors.append(f'{fname} {lname}')
                db.session.commit()

            if 'first_name_passenger_4' in request.form and flight_id:
                flight = db.session.get(Flight,flight_id)
                flight.first_name_passenger_4 = request.form.get('first_name_passenger_4')
                flight.last_name_passenger_4 = request.form.get('last_name_passenger_4')
                flight.passenger_4_type = request.form.get('passenger_4_type')
                fname = request.form.get('first_name_passenger_4')
                lname = request.form.get('last_name_passenger_4')
                if 'passenger_4_type' in request.form:
                    minors.append(f'{fname} {lname}')
                else:
                    majors.append(f'{fname} {lname}')
                db.session.commit()

            if 'first_name_passenger_5' in request.form and flight_id:
                flight = db.session.get(Flight,flight_id)
                flight.first_name_passenger_5 = request.form.get('first_name_passenger_5')
                flight.last_name_passenger_5 = request.form.get('last_name_passenger_5')
                flight.passenger_5_type = request.form.get('passenger_5_type')
                fname = request.form.get('first_name_passenger_5')
                lname = request.form.get('last_name_passenger_5')
                if 'passenger_5_type' in request.form:
                    if 'passenger_4_type' in request.form:
                        minors.append(f'{fname} {lname}')
                    else:
                        majors.append(f'{fname} {lname}')
                db.session.commit()
                
            if 'first_name' in request.form and flight_id:
                flight = db.session.get(Flight,flight_id)
                flight.last_name = request.form.get('last_name')
                flight.first_name = request.form.get('first_name')
                flight.email = request.form.get('email')
                db.session.commit()
            

            if 'address_line1' in request.form and flight_id:
                flight = db.session.get(Flight,flight_id)
                flight.address_line1 = request.form.get('address_line1')
                flight.address_line2 = request.form.get('address_line_2')
                flight.city = request.form.get('city')
                flight.postal_code = request.form.get('postal_code')
                flight.country = request.form.get('country')
                flight.phone_number = request.form.get('phone_number')
                flight.notification = request.form.get('notification')
                db.session.commit()
                print("Flight updated with address:", flight)

            if 'contacted_airline' in request.form and flight_id:
                flight = db.session.get(Flight,flight_id)
                flight.contacted_airline = request.form.get('contacted_airline')
                flight.incident_description = request.form.get('incident_description')
                db.session.commit()

            if 'signature' in request.form and flight_id:
                print("cccccccccccccccccc")
                signature = request.form.get('signature')
                if signature:
                    flight = db.session.get(Flight,flight_id)
                    # flight.signature = 'Yes'
                    header, encoded = signature.split(",")
                    image_data = base64.b64decode(encoded)
                    image = Image.open(BytesIO(image_data))

                    signature_file_path =  os.path.join(app.config['UPLOAD_FOLDER'], f'{user.id}/signature_{flight_id}.png')
                    image.save(signature_file_path)
                    # from consent import consent_form
                    # consent_form(user.id,user.first_name, flight.flight_number, flight.airline , encoded,flight_id)
                    guardian_name = f'{user.first_name} {user.last_name}'
                    
                    flight = db.session.get(Flight,flight_id)
                    from time import sleep
                    if flight.group_travel=='yes':
                        if minors:
                            relation_name = ''
                            for index, minor in enumerate(minors):
                                relation_name+=f'{minor}, '
                            from consent_form import consent_form_minor
                            consent_form_minor(user.id,flight_id,guardian_name, relation_name, flight.flight_number,flight.airline)
                            sleep(2)
                        if majors:
                            relation_name = ''
                            for index, major in enumerate(majors):
                                relation_name+=f'{major}, '
                            from consent_form import consent_form_major
                            consent_form_major(user.id,flight_id,guardian_name, relation_name, flight.flight_number,flight.airline)
                            sleep(2)
                    else:
                        from consent_form import consent_form_individual
                        consent_form_individual(user.id,flight_id,guardian_name, flight.flight_number,flight.airline)
                        sleep(2)

                    forms = [form for form in os.listdir(f'uploads/{user.id}') if form.endswith(f'{flight_id}.pdf')]
                    url = f"{url_for('forget_password', _external=True)}"
                    try:
                        print(first_time_user)
                        if first_time_user:
                            print("FURST")
                            from mailer import welcome_mail
                            password = session['password']
                            welcome_mail(user.id, user.first_name, user.email, password, url, flight_id, forms)
                            first_time_user = False
                            session.pop('password')
                        else:
                            print("Last")
                            from mailer import application_submitted
                            application_submitted(user.id, user.first_name, user.email, url, flight_id, forms)

                            # db.session.commit()
                            # print(majors, minors)
                        
                    except Exception as e:
                        print(f"An error occurred: {e}")
                            

                   
                    
                    

                    # with open(signature_file_path, 'rb') as file:
                    #     signature_binary = file.read()

                    # flight.signature = signature_binary
                    # db.session.commit()
            
                    # flash('Your claim has been successfully submitted!')
                    # return redirect(url_for('login'))
                else:
                    flash('Signature is empty.')
            else:
                print("not working")

            if 'documents' in request.files and flight_id:
                files = request.files.getlist('documents')
                for index, file in enumerate(files):
                    if file and allowed_file(file.filename):
                        filename = secure_filename(str(index)+str(flight_id)+'_'+file.filename)
                        file_path = os.path.join(app.config['UPLOAD_FOLDER'], f'{user.id}/'+filename)
                        file.save(file_path)
                flash('Your claim information has been saved successfully!')
                return redirect(url_for('login'))

        except Exception as e:
            print(f"An error occurred: {e}")
            flash('An error occurred while saving data.')
            return redirect(url_for('login'))
           
    state='login'
    if 'role' in session.keys():
        state='Dashboard'
    data = dict()
    data['first_name'] = ''
    data['last_name'] = ''
    data['leaving_from'] = ''
    data['destination'] = ''
    data['airline'] = ''
    return render_template('form.html',state=state,data=data)

@app.route('/index1', methods=['GET', 'POST'])
def index1():
    #form data
    if request.method == 'POST':
        first_time_user = False
        if 'user_id' not in session:
            first_name = request.form.get('first_name','John')
            last_name = request.form.get('last_name','Doe')
            email = request.form.get('email').strip().lower()
            if email:
                if User.query.filter_by(email=email).first():
                    flash('Your email is already registerd, login please!')
                    return redirect(url_for('login'))
        
            password = gen_passwd()
            user_id = gen_uid(User)
            # user.id = user_id
            session['user_id'] = user_id
            session['username'] = f"{first_name} {last_name}"
            session['role'] = 'user'
            session['email'] = email
            session['password'] = password
            user = User(id=user_id,first_name=first_name, last_name=last_name, email=email, password=generate_password_hash(password),role='user')
            db.session.add(user)
            db.session.commit()
            # db.session.commit()
            first_time_user = True

        user = db.session.get(User, session['user_id']) 
        try:
            if 'leaving_from' in request.form:
                flight = Flight(
                    id = gen_uid(Flight),
                    user_id=user.id,
                    leaving_from=request.form.get('leaving_from'),
                    destination=request.form.get('destination'),
                    connecting_flights=request.form.get('connecting_flights')
                )
                db.session.add(flight)
                db.session.commit()

                # flight.id = gen_uid(Flight)
                # db.session.commit()
                session['flight_id'] = flight.id
                print("Flight added:", flight)
                print(request.form.get('leaving_from'),request.form.get('destination'))
            flight_id = session.get('flight_id')
            if 'date' in request.form and flight_id:
                flight = db.session.get(Flight,flight_id)
                from datetime import date
                date = str(date.today())
                flight.date = date
                db.session.commit()
            
            if 'disruption' in request.form and flight_id:
                flight = db.session.get(Flight,flight_id)
                flight.disruption = request.form.get('disruption')  # Assuming this is what you meant
                flight.delay = request.form.get('delay')  # Assuming this is what you meant
                db.session.commit()
                print("Flight updated with disruption:", flight)

            if 'airline' in request.form and flight_id:
                flight = db.session.get(Flight,flight_id)
                flight.airline = request.form.get('airline')
                flight.flight_number = request.form.get('flight_number')
                flight.flight_date = request.form.get('flight_date')
                db.session.commit()
            
        
            if 'group_travel' in request.form and flight_id:
                flight = db.session.get(Flight,flight_id)
                flight.group_travel = request.form.get('group_travel')
                db.session.commit()
        
            if 'first_name_passenger_2' in request.form and flight_id:
                flight = db.session.get(Flight,flight_id)
                flight.first_name_passenger_2 = request.form.get('first_name_passenger_2')
                flight.last_name_passenger_2 = request.form.get('last_name_passenger_2')
                flight.email_passenger_2= request.form.get('email_passenger_2')
                db.session.commit()
            
            
            if 'first_name' in request.form and flight_id:
                flight = db.session.get(Flight,flight_id)
                flight.last_name = request.form.get('last_name')
                flight.first_name = request.form.get('first_name')
                flight.email = request.form.get('email')
                db.session.commit()
        

            if 'address_line1' in request.form and flight_id:
                flight = db.session.get(Flight,flight_id)
                flight.address_line1 = request.form.get('address_line1')
                flight.address_line2 = request.form.get('address_line_2')
                flight.city = request.form.get('city')
                flight.postal_code = request.form.get('postal_code')
                flight.state = request.form.get('state')
                flight.country = request.form.get('country')
                flight.phone_number = request.form.get('phone_number')
                flight.email = request.form.get('email')
                db.session.commit()
                print("Flight updated with address:", flight)

            if 'booking_number' in request.form and flight_id:
                flight = db.session.get(Flight,flight_id)
                flight.booking_number = request.form.get('booking_number')
                db.session.commit()
                print("Flight updated with booking number:", flight)

            if 'signature' in request.form and flight_id:
                signature = request.form.get('signature')
                if signature:
                    flight = db.session.get(Flight,flight_id)
                    header, encoded = signature.split(",", 1)
                    from consent import consent_form
                    from mailer import welcome_mail
                    consent_form(user.id,user.first_name, flight.flight_number, flight.airline , encoded,flight_id)
                    if first_time_user:
                        url = f"{url_for('forget_password', _external=True)}"
                        password = session['password']
                        welcome_mail(user.id,user.first_name,user.email,password,url,flight_id)
                        first_time_user=False
                        session.pop('password')
                else:
                    flash('Signature is empty.')

            if 'documents' in request.files and flight_id:
                files = request.files.getlist('documents')
                print(files)
                for index, file in enumerate(files):
                    if file and allowed_file(file.filename):
                        filename = secure_filename(str(index)+file.filename)
                        file_path = os.path.join(app.config['UPLOAD_FOLDER'], f'{user_id}/'+filename)
                        file.save(file_path)
                flash('Your claim information has been saved successfully!')
                return redirect(url_for('login'))

        except Exception as e:
            print(f"An error occurred: {e}")
            flash('An error occurred while saving data.')
    state='login'
    if 'role' in session.keys():
        state='Dashboard'
    data = dict()
    data['first_name'] = ''
    data['last_name'] = ''
    data['leaving_from'] = ''
    data['destination'] = ''
    data['airline'] = ''
    return render_template('form.html', state='login', data=data)

@app.route('/adminlogin', methods=['GET', 'POST'])
def admin_login():
    if request.method == 'POST':
        email = request.form.get('email').strip().lower()
        password = request.form.get('password')

        user = db.session.query(User).filter_by(email=email).first()

        if user and user.role in ['admin', 'super_admin']:
            if check_password_hash(user.password, password):
                session['user_id'] = user.id
                session['username'] = f"{user.first_name} {user.last_name}"
                session['email'] = email
                session['role'] = user.role

                # Set admin-related session data
                if user.role == 'admin':
                    session['admin_id'] = user.id
                    session['admin_name'] = session['username']
                    session['role'] = 'partner'  # 👈 Treat as 'partner' for dashboard compatibility
                    return redirect(url_for('partner_dashboard'))

                elif user.role == 'super_admin':
                    return redirect(url_for('super_dashboard'))

                else:
                    flash('Invalid role assigned.')
                    return redirect(url_for('admin_login'))
            else:
                flash('Incorrect password.')
                return redirect(url_for('admin_login'))
        else:
            flash('Invalid email or unauthorized access.')
    return render_template('admin_login.html')

@app.route('/admin_dashboard', methods=['GET','POST'])
def admin_dashboard():
    if 'user_id' not in session or session['role']!='admin':
        flash('You need to log in first.')
        return redirect(url_for('admin_login'))
    user_id = session['user_id']
    role=session['role']
    username = session['username']
    # dates = [date[0] for date in db.session.query(Flight.date).all()]


    PER_PAGE =10 
    # Retrieve filter values from the form (GET request)
    date = request.args.get('date')
    admin = request.args.get('admin')
    status = request.args.get('status')
    
    query = db.session.query(Flight)
    query = query.join(User, Flight.user_id == User.id).filter(User.assigned_admin_id == user_id)
    dates = [date[0] for date in query.with_entities(Flight.date).all()]
    status_all =[status[0] for status in query.with_entities(Flight.status).all()]
    status_all = list(set(status_all))
    # Retrieve the current page number, defaulting to 1
    page = request.args.get('page', 1, type=int)

    # Apply filters dynamically if they are provided
    if date:
        query = query.filter(Flight.date==date)  # Case-insensitive match
    
    if status:
        query = query.filter(Flight.status==status)  # Case-insensitive match
    
    # Paginate the results, specifying the number of results per page (PER_PAGE)
    pagination = query.paginate(page=page, per_page=PER_PAGE, error_out=False)
    
    # Fetch the current page of results
    flights = pagination.items
    
    # Render the template with the filtered results and pagination data
    return render_template("admin_dashboard.html", flights=flights,role=role,username=username,pagination=pagination,allowed_dates=dates,allowed_status=status_all,date=date,status=status)

@app.route('/update_flight_status/<int:flight_id>', methods=['POST'])
def update_flight_status(flight_id):
    if 'user_id' not in session or session.get('role') != 'admin':
        flash('You do not have access to this page.')
        return redirect(url_for('admin_login'))

    flight = db.session.get(Flight,flight_id)
    if flight:
        flight.status = request.form.get('status')
        db.session.commit()
        flash('Flight status updated successfully.')
    else:
        flash('Flight not found.')

    return redirect(url_for('admin_dashboard'))

@app.route('/create_admin', methods=['GET', 'POST'])
def create_admin():
    if 'user_id' not in session or session.get('role') != 'super_admin':
        flash('You do not have access to this page.')
        return redirect(url_for('login'))

    if request.method == 'POST':
        print('post')
        first_name = request.form.get('name')
        # last_name = 'admin'
        last_name = request.form.get('last_name', 'admin')  # Allow setting last name
        email = request.form.get('email').strip().lower()
        password = request.form.get('password')
        role = request.form.get('role','admin')

        if db.session.query(User).filter_by(email=email).first():
            flash('Email address already registered. Please use a different email.')
            print('not ok')
        else:
            print('ok')
            new_admin = User(first_name=first_name, last_name=last_name, email=email, password=generate_password_hash(password), role=role)
            db.session.add(new_admin)
            db.session.commit()
            flash('New admin created successfully!')
            return redirect(url_for('create_admin'))
    admins = db.session.query(User).filter(User.role.in_(['admin','super_admin','partner'])).all()
    return render_template('addadmin.html', admins=admins)

@app.route('/super-dash', methods=['GET', 'POST'])
def super_dashboard():
    if 'user_id' not in session or session['role']!='super_admin':
        flash('You need to log in first.')
        return redirect(url_for('admin_login'))
    user_id = session['user_id']
    role=session['role']
    username = session['username']
    admins = db.session.query(User).filter_by(role='admin').all()
    flights = db.session.query(Flight).all()
    status_all =[status[0] for status in db.session.query(Flight.status).all()]
    status_all = list(set(status_all))
    dates = [date[0] for date in db.session.query(Flight.flight_date).all()]


    query = db.session.query(Flight)
    PER_PAGE =10
    # Retrieve filter values from the form (GET request)
    date = request.args.get('date')
    admin = request.args.get('admin')
    status = request.args.get('status')
    
    # Retrieve the current page number, defaulting to 1
    page = request.args.get('page', 1, type=int)

    # Apply filters dynamically if they are provided
    if date:
        query = query.filter(Flight.date==date)  # Case-insensitive match
    
    if admin:
        query = query.join(User, Flight.user_id == User.id).filter(User.assigned_admin_id == admin)
    
    if status:
        query = query.filter(Flight.status==status)  # Case-insensitive match
    
    # Paginate the results, specifying the number of results per page (PER_PAGE)
    pagination = query.paginate(page=page, per_page=PER_PAGE, error_out=False)
    
    # Fetch the current page of results
    flights = pagination.items
    
    # Render the template with the filtered results and pagination data
    return render_template("anly.html", flights=flights,role=role,username=username,admins=admins,pagination=pagination,allowed_dates=dates,allowed_status=status_all,date=date,admin=admin,status=status)

 
@app.route('/super', methods=['GET','POST'])
def super():
    if 'user_id' not in session or session['role']!='super_admin':
        flash('You need to log in first.')
        return redirect(url_for('admin_login'))
    if request.method == 'POST':
        status = request.form.get('option')
        id = request.form.get('id')
        if status and id:
            flight = db.session.query(Flight).filter_by(id=id).first()
            flight.status = status
            db.session.commit()
    user_id = session['user_id']
    role=session['role']
    username = session['username']
    admins = db.session.query(User).filter_by(role='admin').all()
    flights = db.session.query(Flight).all()
    status_all =[status[0] for status in db.session.query(Flight.status).all()]
    status_all = list(set(status_all))
    dates = [date[0] for date in db.session.query(Flight.flight_date).all()]
    dates = list(set(dates))

    query = db.session.query(Flight)
    PER_PAGE =10
    # Retrieve filter values from the form (GET request)
    date = request.args.get('date')
    admin = request.args.get('admin')
    status = request.args.get('status')
    
    # Retrieve the current page number, defaulting to 1
    page = request.args.get('page', 1, type=int)

    # Apply filters dynamically if they are provided
    if date:
        query = query.filter(Flight.flight_date==date)  # Case-insensitive match
    
    if admin:
        query = query.join(User, Flight.user_id == User.id).filter(User.assigned_admin_id == admin)
    
    if status:
        query = query.filter(Flight.status==status)  # Case-insensitive match
    
    # Paginate the results, specifying the number of results per page (PER_PAGE)
    pagination = query.paginate(page=page, per_page=PER_PAGE, error_out=False)
    
    # Fetch the current page of results
    flights = pagination.items
    
    # Render the template with the filtered results and pagination data
    return render_template("crm.html", flights=flights,role=role,username=username,admins=admins,pagination=pagination,allowed_dates=dates,allowed_status=status_all,date=date,admin=admin,status=status)


@app.route('/admin_crm', methods=['GET','POST'])
def admin_crm():
    if 'user_id' not in session or session['role']!='admin':
        flash('You need to log in first.')
        return redirect(url_for('admin_login'))
    if request.method == 'POST':
        status = request.form.get('option')
        id = request.form.get('id')
        if status and id:
            flight = db.session.query(Flight).filter_by(id=id).first()
            flight.status = status
            db.session.commit()
    user_id = session['user_id']
    role=session['role']
    username = session['username']

    #base query
    query = db.session.query(Flight)
    query = query.join(User, Flight.user_id == User.id).filter(User.assigned_admin_id == user_id)
    dates = [date[0] for date in query.with_entities(Flight.date).all()]
    dates = list(set(dates))
    status_all =[status[0] for status in query.with_entities(Flight.status).all()]
    status_all = list(set(status_all))

    PER_PAGE =10
    # Retrieve filter values from the form (GET request)
    date = request.args.get('date')
    status = request.args.get('status')
    
    # Retrieve the current page number, defaulting to 1
    page = request.args.get('page', 1, type=int)

    # Apply filters dynamically if they are provided
    if date:
        query = query.filter(Flight.flight_date==date)  # Case-insensitive match
    
    if status:
        query = query.filter(Flight.status==status)  # Case-insensitive match
    
    # Paginate the results, specifying the number of results per page (PER_PAGE)
    pagination = query.paginate(page=page, per_page=PER_PAGE, error_out=False)
    
    # Fetch the current page of results
    flights = pagination.items
    
    # Render the template with the filtered results and pagination data
    return render_template("admin_crm.html", flights=flights,role=role,username=username,pagination=pagination,allowed_dates=dates,allowed_status=status_all,date=date,status=status,admin=user_id)


@app.route('/view/<int:flight_id>', methods=['GET','POST'])
def view(flight_id):
    if 'user_id' not in session or session.get('role') not in ['super_admin','admin']:
        flash('You do not have access to this page.')
        return redirect(url_for('login'))
    if request.method == 'POST':
        if request.form.get('comment'):
            flight = db.session.query(Flight).filter_by(id=flight_id).one()
            comment = request.form.get('comment')
            admin_name=session['username']
            new_comment = Comment(text=comment,admin_name=admin_name,user_id=session['user_id'],flight_id=flight_id)

            # Add the new Comment object to the InstrumentedList
            flight.comments.append(new_comment)
            db.session.commit()
    
    flight = db.session.query(Flight).filter_by(id=flight_id).one()
    return render_template('view.html',flight=flight)

@app.route('/view_docs/<int:user_id>/<int:flight_id>', methods=['GET'])
def view_docs(user_id, flight_id):
    if 'user_id' not in session or session.get('role') not in ['super_admin', 'admin']:
        flash('You do not have access to this page.')
        return redirect(url_for('login'))

    folder_path = os.path.join(app.config['UPLOAD_FOLDER'], f'partner_client_{flight_id}')
    files = []

    if os.path.exists(folder_path):
        try:
            files = [file for file in os.listdir(folder_path) if os.path.isfile(os.path.join(folder_path, file))]
        except Exception as e:
            flash(f'Error reading uploaded files: {e}')
            return redirect(url_for('partner_dashboard'))
    else:
        flash('No uploaded documents found.')

    flight = db.session.get(Flight, flight_id)
    signed_file = os.path.basename(flight.signed_file_path) if flight and flight.signed_file_path else None

    return render_template(
        'view_docs.html',
        files=files,
        signed_file=signed_file,
        id=str(user_id),
        flight_id=flight_id
    )


@app.route('/download_doc/<int:flight_id>/<filename>')
def download_doc(flight_id, filename):
    path = os.path.join(app.config['UPLOAD_FOLDER'], f'partner_client_{flight_id}')
    return send_from_directory(path, filename, as_attachment=False)

@app.route('/doc',methods=['GET','POST'])
def doc():
    id = request.args.get('id')
    filename = request.args.get('filename')
    # path = os.path.join('uploads','test_folder')
    path = f'uploads/{str(id)}'
    return send_from_directory(path,filename)

def get_user_flights(user_id):
    return Flight.query.filter_by(user_id=user_id).all()
@app.route('/user_dashboard')
def user_dashboard():
    if 'user_id' not in session:
        return redirect(url_for('login'))

    user_id = session['user_id']
    flights = get_user_flights(user_id)
    username = session.get('username')
    icon = username[0] if username else '?'
    
    return render_template(
        'user2dash.html',
        flights=flights,
        username=username,
        icon=icon
    )




@app.route('/in_dash/<int:flight_id>', methods=['GET', 'POST'])
def in_dash(flight_id):
    if 'role' not in session and session['role']!='user':
        return redirect(url_for('login'))
    msg = request.args.get('msg','')
    if request.method=='POST':
        try:
            user_id = session['user_id']
            if 'files[]' in request.files and flight_id:
                    files = request.files.getlist('files[]')
                    n = 0
                    for index, file in enumerate(files):
                        if file and allowed_file(file.filename):
                            filename = secure_filename(str(index+1)+"_new_"+file.filename)
                            file_path = os.path.join(app.config['UPLOAD_FOLDER'], f'{user_id}/'+filename)
                            file.save(file_path)
                            n = index+1
                     
                    msg = f'{n} file uploaded successfully!' if n>0 else 'file type not supported!'
                    session['msg']=msg
                    flight = db.session.query(Flight).filter_by(id=flight_id).one()
                    return redirect(url_for('in_dash', flight_id=flight_id,msg=msg))
            else:
                session['msg']=msg
                msg = 'file not selected!'
        except Exception as e:
            session['msg']=msg
            msg = 'an error occured!'
    msg = session['msg'] if 'msg' in session.keys() else ''
    if 'msg' in session.keys():
        del session['msg'] 
    flight = db.session.query(Flight).filter_by(id=flight_id).one()
    return render_template('claim_submit_next_user.html', flight=flight,msg=msg)

@app.route('/settings', methods=['GET','POST'])
def settings():
    if 'user_id' not in session.keys():
        return redirect(url_for('login'))
    
    user_id = session['user_id']
    username = session['username']
    icon = list(username.split()[0])[0] + list(username.split()[1])[0]
    
    user = db.session.query(User).filter_by(id=user_id).one()
    if request.method=='POST':
        type = request.args.get('type',None)
        if type == 'details':
            if 'first_name' in request.form and user:
                user.first_name = request.form.get('first_name')
                db.session.commit()
                session['username'] = user.first_name

            if 'last_name' in request.form and user:
                user.last_name = request.form.get('last_name')
                db.session.commit()
                session['username']+= " " + user.first_name

            if 'email' in request.form and user:
                user.first_name = request.form.get('email')
                db.session.commit()
        elif type=='password':
            old = request.form.get('old')
            new = request.form.get('new')
            confirm = request.form.get('confirm')
            if check_password_hash(user.password,old):
                if new == confirm:
                    user.password = generate_password_hash(new)
                    flash('password has been updated!')
                    db.session.commit()
                else:
                    flash('new password and confirm password did not match!')
            else:
                flash(user.password)
                flash(generate_password_hash(old))
                flash('wrong old password!')
    data = {'first_name':user.first_name, 'last_name':user.last_name, 'email':user.email}
    return render_template('User_settings.html',data=data,icon=icon)


@app.route('/delete/<int:flight_id>')#, methods=['POST'])
def delete(flight_id):
    flight = db.session.get(Flight,flight_id)
    if flight:
        db.session.delete(flight)
        db.session.commit()
        flash('Flight deleted successfully.')
    else:
        flash('Flight not found.')

    return redirect(url_for('dashboard'))


@app.route('/add_comment/<int:user_id>', methods=['POST'])
def add_comment(user_id):
    if 'user_id' not in session:
        flash('You need to log in first.')
        return redirect(url_for('login'))

    comment_text = request.form.get('comment')

    if comment_text:
        new_comment = Comment(user_id=user_id, text=comment_text)
        db.session.add(new_comment)
        db.session.commit()
        flash('Comment added successfully!')
    else:
        flash('Comment cannot be empty.')

    return redirect(url_for('admin_dashboard'))


@app.route('/login', methods=['GET', 'POST'])
def login():
    # If POST, attempt login
    if request.method == 'POST':
        # Retrieve email and password from form
        email = request.form.get('email', '').strip().lower()
        password = request.form.get('password', '')
        user = db.session.query(User).filter_by(email=email).first()
        if user and check_password_hash(user.password, password):
            session['user_id'] = user.id
            session['username'] = user.first_name + " " + user.last_name
            session['role'] = user.role
            session['email'] = user.email
            # Correct redirect after login
            if user.role == 'admin':
                return redirect(url_for('admin_dashboard'))
            elif user.role == 'partner':
                return redirect(url_for('partner_dashboard'))
            else:
                return redirect(url_for('user_dashboard'))

        flash('Invalid email or password.')

    # If already logged in, redirect out of login
    elif 'user_id' in session:
        role = session.get('role')
        if role == 'admin':
            return redirect(url_for('admin_dashboard'))
        elif role == 'partner':
            return redirect(url_for('partner_dashboard'))
        else:
            return redirect(url_for('user_dashboard'))

    return render_template('login.html')


@app.route('/logout')
def logout():
    session.pop('user_id', None)
    session.pop('username', None)
    session.pop('role', None)
    session.clear()
    flash('You have been logged out.', 'success')
    return redirect(url_for('login'))


@app.route('/update_claims_status', methods=['POST'])
def update_claims_status():
    for claim_id in request.form:
        if claim_id.startswith('status_'):
            claim_id_clean = claim_id.split('_')[1]
            new_status = request.form.get(claim_id)
            # Assuming you have a Claim model
            claim = db.session.get(Flight,claim_id_clean)
            if claim:
                claim.status = new_status
                db.session.commit()
    flash('Claim statuses updated successfully!')
    return redirect(url_for('admin_dashboard'))

@app.route('/assign_users', methods=['GET', 'POST'])
def assign_users():
    if 'user_id' not in session or session.get('role') != 'super_admin':
        flash('You do not have access to this page.')
        return redirect(url_for('admin_login'))

    admin_id = request.form.get('admin_id')
    user_ids = request.form.getlist('user_ids[]')

    if admin_id and user_ids:
        for user_id in user_ids:
            user = db.session.get(User,user_id)
            if user:
                user.assigned_admin_id = admin_id
        db.session.commit()
        flash('Users successfully assigned to the selected admin.')

    return redirect(url_for('super_dashboard'))
@app.route('/forget_password', methods=['GET', 'POST'])
def forget_password():
    if request.method=='POST':
        email = request.form.get('email')
        user = db.session.query(User).filter_by(email=email).first()
        token = request.args.get('token')
        if user:
            from mailer import pass_reset
            import uuid  # To generate a unique verification token
            username = user.first_name + user.last_name
            email = user.email
            # Generate a unique verification token
            token = str(uuid.uuid4())
            url = f"{url_for('reset_password',id=user.id, token=token, _external=True)}"
            print(url)
            #send time
            expiry_time = time.perf_counter() + 10*60
            #add these to session for verification 
            session['token'] = token
            session['expiry_time'] = expiry_time
            print(email, username,url)
            pass_reset(email, username,url)
            msg='password reset url sent to your registered email!'
            flash(msg)
            return render_template('login.html')
        else:
            msg = 'invalid email!'
            flash(msg)
            return render_template('forget_password.html')
    else:
        return render_template('forget_password.html')
    
@app.route('/reset_password', methods=['POST','GET'])
def reset_password():
    id = request.args.get('id')
    if request.method=='POST':
        new_password = request.form.get('new_password')
        confirm_new_password = request.form.get('confirm_new_password')
        if id:
            user = db.session.query(User).filter_by(id=id).first()
            if new_password==confirm_new_password:
                user.password = generate_password_hash(new_password)
                db.session.commit()
                flash('password has been updated!')
                return redirect(url_for('login'))
            else:
                flash('new password did not match confirm password!')
                return redirect(url_for('reset_password'))
        else:
            flash('invalid url!')
            return redirect(url_for('forget_password'))
        
    token = request.args.get('token')
    if token:
        if token in session.values():
            # print(token)
            if session['expiry_time']>=time.perf_counter():
                return render_template('confirm_password.html',id=id)
            else:
                session.pop(token)
                del session['expiry_time']
                msg = 'url expired!'
                flash(msg)
                return render_template('forget_password.html')
        else:
            msg = 'invalid url!'
            flash(msg)
            return render_template('forget_password.html')
    else:
        msg = 'invalid url!'
        flash(msg)
        return render_template('forget_password.html')
        

@app.route('/suggest', methods=['GET'])
def suggest():
    query = request.args.get('q', '').lower()
    suggestions = filter_airports(query)
    # results = [s for s in suggestions if s.startswith(query)]
    return jsonify(suggestions)

@app.route('/suggest_airlines', methods=['GET'])
def suggest_airlines():
    query = request.args.get('q', '').lower()
    suggestions = filter_airlines(query)
    # results = [s for s in suggestions if s.startswith(query)]
    return jsonify(suggestions)

@app.route('/partner_login', methods=['GET', 'POST'])
def partner_login():
    if request.method == 'POST':
        email = request.form.get('email').strip().lower()
        password = request.form.get('password')

        # Query the database using SQLAlchemy
        user = db.session.query(User).filter_by(email=email).first()

        if user and user.role == 'partner':  # Ensure only partners can log in here
            if check_password_hash(user.password, password):
                session['user_id'] = user.id
                session['username'] = f"{user.first_name} {user.last_name}"
                session['role'] = user.role
                session['email'] = email
                return redirect(url_for('partner_dashboard'))
            else:
                flash('Invalid password!')
                return redirect(url_for('partner_login'))
        else:
            flash('Invalid email or unauthorized access!')
    return render_template('partner_login.html')

@app.route('/partner_dashboard', methods=['GET', 'POST'])
def partner_dashboard():
    if 'user_id' not in session or session['role'] != 'partner':
        flash('You need to log in first.')
        return redirect(url_for('partner_login'))
        
    user_id = session['user_id']
    role = session['role']
    username = session['username']
    
    # Get filter parameters
    date = request.args.get('date', '')  # Get date parameter, default to empty string
    status = request.args.get('status', '')  # Get status parameter, default to empty string
    page = request.args.get('page', 1, type=int)
    
    # Debug information
    print(f"Received filter params - date: '{date}', status: '{status}'")
    
    # Start the query - show all data for this partner by default
    query = Flight.query.filter_by(partner_id=user_id)

    # Apply filters only if they are provided and not empty
    if date and date.strip():
        print(f"Applying date filter: '{date}'")
        # Try both date fields since we're not sure which one is used
        date_query = query.filter(
            db.or_(
                Flight.flight_date == date,
                Flight.date == date
            )
        )
        query = date_query
            
    if status and status.strip():
        print(f"Applying status filter: '{status}'")
        query = query.filter(Flight.status == status)
    
    # Pagination
    pagination = query.paginate(page=page, per_page=10, error_out=False)
    flights = pagination.items
    
    # Get unique dates for filter dropdowns
    # Try both date fields
    flight_dates = db.session.query(Flight.flight_date).filter(
        Flight.partner_id == user_id,
        Flight.flight_date.isnot(None)
    ).distinct().all()
    
    regular_dates = db.session.query(Flight.date).filter(
        Flight.partner_id == user_id,
        Flight.date.isnot(None)
    ).distinct().all()
    
    # Combine and clean dates
    all_dates = flight_dates + regular_dates
    allowed_dates = [str(date[0]) for date in all_dates if date[0]]
    allowed_dates = sorted(list(set(allowed_dates)))  # Remove duplicates and sort
    
    # Debug: Print the allowed dates
    print(f"Available dates for filtering: {allowed_dates}")

    # Update allowed status list to include all required statuses
    allowed_status = ['open', 'progress', 'pending', 'waiting_signature', 'forwarded', 'closed']
    
    # Count flights by status
    open_count = Flight.query.filter_by(partner_id=user_id, status='open').count()
    progress_count = Flight.query.filter_by(partner_id=user_id, status='progress').count()
    pending_count = Flight.query.filter_by(partner_id=user_id, status='pending').count()
    waiting_signature_count = Flight.query.filter_by(partner_id=user_id, status='waiting_signature').count()
    forwarded_count = Flight.query.filter_by(partner_id=user_id, status='forwarded').count()
    closed_count = Flight.query.filter_by(partner_id=user_id, status='closed').count()
    
    # Debug information
    print(f"Filter params - date: '{date}', status: '{status}'")
    print(f"Status counts - open: {open_count}, progress: {progress_count}, pending: {pending_count}, " 
          f"waiting_signature: {waiting_signature_count}, forwarded: {forwarded_count}, closed: {closed_count}")
    print(f"Total flights after filtering: {pagination.total}")
    
    return render_template("partner_dashboard.html", 
                          role=role, 
                          username=username, 
                          flights=flights,
                          pagination=pagination,
                          allowed_dates=allowed_dates,
                          allowed_status=allowed_status,
                          open_count=open_count,
                          progress_count=progress_count,
                          pending_count=pending_count,
                          waiting_signature_count=waiting_signature_count,
                          forwarded_count=forwarded_count,
                          closed_count=closed_count,
                          date=date,
                          status=status)
# @app.route('/partner')
# def partner():
#     # Check if user is already logged in as a partner
#     if 'role' in session and session['role'] == 'partner':
#         return redirect(url_for('partner_dashboard'))
        
#     return render_template('partner_login.html')

# Add this function to your app.py
def allowed_file(filename, allowed_extensions=None):
    if allowed_extensions is None:
        allowed_extensions = app.config['ALLOWED_EXTENSIONS']
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in allowed_extensions
@app.route('/partner_upload', methods=['GET', 'POST'])
def partner_upload():
    if 'user_id' not in session or session['role'] != 'partner':
        flash('You need to log in first.')
        return redirect(url_for('partner_login'))

    user_id = session['user_id']
    role = session['role']
    username = session['username']
    message = None

    if request.method == 'POST':
        file = request.files.get('file')
        if not file or file.filename == '':
            message = 'No file selected.'
            return render_template('partner_upload.html', role=role, username=username, message=message)

        if file and allowed_file(file.filename, {'csv'}):
            import tempfile, chardet, csv, secrets, random, string, datetime
            from io import StringIO
            from docx import Document
            from mailer import partner_client_email  # ✅ Use partner_client_email only

            def generate_password(length=10):
                chars = string.ascii_letters + string.digits + "!@#$%="
                return ''.join(random.choices(chars, k=length))

            def generate_reference_id():
                last_flight = Flight.query.order_by(Flight.reference_id.desc()).first()
                return last_flight.reference_id + 1 if last_flight and last_flight.reference_id else 10000

            try:
                with tempfile.NamedTemporaryFile(delete=False, suffix='.csv') as temp_file:
                    file.save(temp_file.name)
                    temp_path = temp_file.name

                with open(temp_path, 'rb') as f:
                    raw_data = f.read()
                    detected_encoding = chardet.detect(raw_data)['encoding']

                encodings_to_try = [detected_encoding, 'utf-8', 'latin-1', 'iso-8859-1', 'cp1252']
                csv_content = None
                for enc in encodings_to_try:
                    try:
                        with open(temp_path, 'r', encoding=enc) as f:
                            csv_content = f.read().replace('\xa0', ' ')
                            successful_encoding = enc
                            break
                    except Exception:
                        continue

                os.unlink(temp_path)
                if not csv_content:
                    message = 'CSV file could not be read. Try UTF-8 or standard Excel export.'
                    return render_template('partner_upload.html', role=role, username=username, message=message)

                stream = StringIO(csv_content)
                csv_reader = csv.DictReader(stream)
                records_added = 0
                new_flight_ids = []

                for row in csv_reader:
                    cleaned_row = {}
                    for key, value in row.items():
                        if key and value:
                            cleaned_key = str(key).strip().replace('\xa0', ' ')
                            cleaned_value = str(value).strip().replace('\xa0', ' ')
                            cleaned_row[cleaned_key] = cleaned_value
                        else:
                            cleaned_row[key] = value

                    token = secrets.token_urlsafe(32)
                    password = generate_password()
                    reference_id = generate_reference_id()

                    new_flight = Flight(
                        first_name=cleaned_row.get('First Name', ''),
                        last_name=cleaned_row.get('Last Name', ''),
                        email=cleaned_row.get('Email ID', ''),
                        phone_number=cleaned_row.get('Phone Number', ''),
                        flight_date=cleaned_row.get('Date', ''),
                        date=cleaned_row.get('Date', ''),
                        airline=cleaned_row.get('Airline', ''),
                        airline2=cleaned_row.get('Connecting Airline', ''),
                        flight_number=cleaned_row.get('Flight Number', ''),
                        leaving_from=cleaned_row.get('Leaving From', ''),
                        destination=cleaned_row.get('Final destination', ''),
                        connecting_flights=cleaned_row.get('connecting flights?', ''),
                        connecting_leaving_from=cleaned_row.get('Connecting Flight', ''),
                        compensation_flight=cleaned_row.get('compenstion for', ''),
                        reason=cleaned_row.get('Delayed due to weather Conditions?', ''),
                        incident_description=cleaned_row.get('what happend', ''),
                        contacted_airline=cleaned_row.get('contacted the airline?', ''),
                        booking_number=cleaned_row.get('Booking Ref No.', ''),
                        address_line1=cleaned_row.get('Address 1', ''),
                        address_line2=cleaned_row.get('Address 2', ''),
                        city=cleaned_row.get('City', ''),
                        postal_code=cleaned_row.get('Postal Code', ''),
                        country=cleaned_row.get('Country', ''),
                        group_travel=cleaned_row.get('traveling with others', ''),
                        partner_id=user_id,
                        status='waiting_signature',
                        verification_token=token,
                        password=password,
                        reference_id=reference_id
                    )

                    db.session.add(new_flight)
                    db.session.flush()
                    new_flight_ids.append(new_flight.id)
                    records_added += 1

                db.session.commit()

                for fid in new_flight_ids:
                    flight = Flight.query.get(fid)

                    template_path = 'consent_form_individual.docx'
                    output_dir = os.path.join(app.config['UPLOAD_FOLDER'], f'partner_docs/{flight.id}')
                    os.makedirs(output_dir, exist_ok=True)
                    output_path = os.path.join(output_dir, f'document_{flight.id}.docx')

                    doc = Document(template_path)
                    for p in doc.paragraphs:
                        if '[guardian_name]' in p.text:
                            p.text = p.text.replace('[guardian_name]', f"{flight.first_name} {flight.last_name}")
                        if '[flight_number]' in p.text:
                            p.text = p.text.replace('[flight_number]', flight.flight_number or '')
                        if '[airline_name]' in p.text:
                            p.text = p.text.replace('[airline_name]', flight.airline or '')
                        if '[date]' in p.text:
                            p.text = p.text.replace('[date]', str(flight.flight_date or datetime.date.today()))

                    doc.save(output_path)
                    flight.signed_file_path = output_path
                    db.session.commit()

                    confirm_url = url_for(
                        'client_complete_claim',
                        flight_id=flight.id,
                        token=flight.verification_token,
                        _external=True
                    )

                    flight_details = f"{flight.airline} {flight.flight_number} on {flight.flight_date}"
                    partner_client_email(
                        client_name=f"{flight.first_name} {flight.last_name}",
                        client_email=flight.email,
                        partner_name=username,
                        flight_details=flight_details,
                        confirm_url=confirm_url
                    )

                message = f"✅ Successfully uploaded {records_added} records using {successful_encoding} encoding."

            except Exception as e:
                db.session.rollback()
                message = f"❌ Error: {e}"

            return render_template('partner_upload.html', role=role, username=username, message=message)

        else:
            message = 'Please upload a valid CSV file.'

    return render_template('partner_upload.html', role=role, username=username, message=message)

def send_client_confirmation_email(flight):
    """Send an email to clients whose claims are submitted by partners"""
    emails_sent = 0  # Initialize emails_sent to avoid NameError
    try:
        import uuid
        from mailer import partner_client_email  # Import the new function
        
        # Generate a unique token for this flight
        token = str(uuid.uuid4())
        
        # Store the token in the database for later verification
        flight.verification_token = token
        db.session.commit()
        
        # Create a secure link that clients can access
        confirm_url = url_for(
            'client_complete_claim',
            flight_id=flight.id,
            token=token,
            _external=True
        )
        
        # Get partner name
        partner = db.session.get(User, flight.partner_id)
        partner_name = f"{partner.first_name} {partner.last_name}" if partner else "Your partner"
        
        # Create flight details string
        flight_details = f"{flight.airline} {flight.flight_number} on {flight.flight_date}"
        
        # Get client name
        client_name = f"{flight.first_name} {flight.last_name}"
        
        # Send email using existing function from mailer.py
        result = partner_client_email(
            client_name=client_name,
            client_email=flight.email,
            partner_name=partner_name,
            flight_details=flight_details,
            confirm_url=confirm_url
        )
        if result:
            emails_sent += 1
        
        return result
        
    except Exception as e:
        print(f"Error in send_client_confirmation_email: {str(e)}")
        return False

@app.route('/complete-claim/<int:flight_id>/<token>', methods=['GET', 'POST'])
def client_complete_claim(flight_id, token):
    flight = db.session.query(Flight).filter_by(id=flight_id, verification_token=token).first()

    if not flight:
        flash('Invalid or expired claim link.')
        return redirect(url_for('landing'))

    if request.method == 'POST':
        try:
            import base64
            from PIL import Image
            from io import BytesIO
            from docx import Document
            from docx.shared import Inches
            import shutil
            from mailer import welcome_mail  # ✅ Make sure mailer is imported

            # Directory to save signature and files
            user_dir_path = os.path.join(os.path.abspath(app.config['UPLOAD_FOLDER']), f"partner_client_{flight.id}")
            os.makedirs(user_dir_path, exist_ok=True)

            # Handle signature
            signature_saved = False
            signature_path = None
            signature_data_url = request.form.get('signature')

            if signature_data_url and signature_data_url.startswith('data:image'):
                try:
                    signature_data = signature_data_url.split(',')[1]
                    signature_bytes = base64.b64decode(signature_data)
                    signature_image = Image.open(BytesIO(signature_bytes))

                    signature_path = os.path.join(user_dir_path, f'signature_{flight.id}.png')
                    signature_image.save(signature_path)
                    signature_saved = True
                    print(f"Signature saved to {signature_path}")
                except Exception as sig_error:
                    print(f"Error saving signature: {sig_error}")

            # Save uploaded documents
            if 'documents' in request.files:
                files = request.files.getlist('documents')
                for index, file in enumerate(files):
                    if file and file.filename and allowed_file(file.filename):
                        try:
                            filename = secure_filename(f"client_{flight.id}_{index}_{file.filename}")
                            file_path = os.path.join(user_dir_path, filename)
                            file.save(file_path)
                        except Exception as file_error:
                            print(f"Error saving file: {file_error}")

            # Update status and regenerate document after signature
            if signature_saved:
                flight.status = 'forwarded'
                doc_path = os.path.join(app.config['UPLOAD_FOLDER'], f'partner_docs/{flight.id}/document_{flight.id}.docx')
                if os.path.exists(doc_path):
                    try:
                        doc = Document(doc_path)
                        doc.add_paragraph("\nSigned electronically by client:")
                        doc.add_picture(signature_path, width=Inches(2.5))
                        doc.save(doc_path)
                        flight.signed_file_path = doc_path
                        print("Consent form updated with signature.")
                    except Exception as doc_error:
                        print(f"Error updating document: {doc_error}")

            db.session.commit()

            # ✅ USER CREATION & WELCOME EMAIL SECTION:
            existing_user = User.query.filter_by(email=flight.email).first()
            if not existing_user:
                try:
                    # store plain password before hashing
                    plain_password = flight.password
                    hashed_password = generate_password_hash(plain_password)

                    new_user = User(
                        email=flight.email,
                        password=hashed_password,
                        first_name=flight.first_name,
                        last_name=flight.last_name,
                        role='user',
                        assigned_admin_id=None
                    )
                    db.session.add(new_user)
                    db.session.commit()

                    # ✅ Link this flight record to the user_id
                    flight.user_id = new_user.id
                    db.session.commit()

                    # Prepare attachment list
                    signed_doc_filename = f'document_{flight.id}.docx'
                    forms = [signed_doc_filename]

                    # Send Welcome Mail (Send plain password, not hashed one)
                    site_url = request.url_root.rstrip('/')
                    welcome_mail(
                        user_id=new_user.id,
                        username=f"{flight.first_name} {flight.last_name}",
                        user_email=new_user.email,
                        password=plain_password,  # ✅ This is correct
                        url=site_url,
                        flight_id=flight.reference_id,
                        forms=forms
                    )

                    print("✅ Login credentials email sent successfully to new CSV user.")

                except Exception as email_error:
                    print(f"⚠️ Error creating user or sending welcome email: {email_error}")


            flash('Thank you! Your claim has been completed and is being processed.')
            return render_template('claim_completed.html', flight=flight)

        except Exception as e:
            import traceback
            print(f"An error occurred in client_complete_claim: {e}")
            print(traceback.format_exc())
            db.session.rollback()
            flash('An error occurred while processing your information. Please try again.')

    return render_template('complete_claim.html', flight=flight)




# @app.route('/add_partner_id_column')
# def add_partner_id_column():
#     try:
#         # Import text from SQLAlchemy
#         from sqlalchemy import text
        
#         # Check if the column already exists - use text() to wrap SQL strings
#         sql_check = text("SELECT COUNT(*) FROM information_schema.COLUMNS WHERE TABLE_SCHEMA = 'defaultdb' AND TABLE_NAME = 'flight' AND COLUMN_NAME = 'partner_id'")
#         result = db.session.execute(sql_check).scalar()
        
#         if result == 0:
#             # Add the column - use text() again
#             sql = text("ALTER TABLE flight ADD COLUMN partner_id INTEGER, ADD CONSTRAINT fk_partner_id FOREIGN KEY (partner_id) REFERENCES user (id)")
#             db.session.execute(sql)
#             db.session.commit()
#             return "Column partner_id added successfully to flight table"
#         else:
#             return "Column partner_id already exists"
#     except Exception as e:
#         db.session.rollback()
#         return f"Error adding column: {str(e)}"
    
@app.route('/check_schema')
def check_schema():
    from sqlalchemy import text
    result = db.session.execute(text("DESCRIBE flight")).fetchall()
    return str(result)

@app.route('/download_signed_doc/<int:flight_id>')
def download_signed_doc(flight_id):
    flight = db.session.get(Flight, flight_id)
    if not flight or not flight.signed_file_path:
        return "Document not found", 404
    directory, filename = os.path.split(flight.signed_file_path)
    return send_from_directory(directory, filename, as_attachment=True)

@app.route('/add_verification_token_column')
def add_verification_token_column():
    try:
        # Import text from SQLAlchemy
        from sqlalchemy import text
        
        # Check if the column already exists
        sql_check = text("SELECT COUNT(*) FROM information_schema.COLUMNS WHERE TABLE_SCHEMA = 'defaultdb' AND TABLE_NAME = 'flight' AND COLUMN_NAME = 'verification_token'")
        result = db.session.execute(sql_check).scalar()
        
        if result == 0:
            # Add the column
            sql = text("ALTER TABLE flight ADD COLUMN verification_token VARCHAR(100)")
            db.session.execute(sql)
            db.session.commit()
            return "Column verification_token added successfully to flight table"
        else:
            return "Column verification_token already exists"
    except Exception as e:
        db.session.rollback()
        return f"Error adding column: {str(e)}"
    
@app.route("/view/<int:id>", methods=["GET", "POST"])
def view_flight(id):
    # Allow partner or admin to view
    if 'user_id' not in session and 'admin_id' not in session:
        flash("Please log in first.")
        return redirect(url_for('login'))

    flight = Flight.query.get_or_404(id)

    if request.method == "POST":
        comment_text = request.form.get("comment", "").strip()
        if comment_text:
            new_comment = Comment(
                text=comment_text,
                flight_id=id,
                admin_name=session.get("admin_name", session.get("username", "Unknown")),
                user_id=session.get("admin_id", session.get("user_id"))
            )
            db.session.add(new_comment)
            db.session.commit()
            flash("Comment added.")

    return render_template("view.html", flight=flight)


@app.route('/download_sample_csv')
def download_sample_csv():
    # Create an in-memory text stream
    csv_data = io.StringIO()
    csv_writer = csv.writer(csv_data)
    
    # Write header row
    csv_writer.writerow([
        'first_name', 'last_name', 'email', 'phone_number', 
        'flight_date', 'airline', 'flight_number', 'address_line1', 
        'city', 'country', 'postal_code', 'status'
    ])
    
    # Write sample data rows
    sample_data = [
        ['John', 'Smith', 'john.smith@example.com', '+1234567890', 
         '2025-06-10', 'British Airways', 'BA123', '123 Main St', 
         'London', 'UK', 'SW1A 1AA', 'open'],
        ['Sarah', 'Johnson', 'sarah.j@example.com', '+1987654321', 
         '2025-06-12', 'Lufthansa', 'LH456', '456 Park Avenue', 
         'Berlin', 'Germany', '10115', 'open'],
        ['Michael', 'Williams', 'm.williams@example.com', '+1122334455', 
         '2025-06-15', 'Emirates', 'EK789', '789 Ocean Blvd', 
         'Dubai', 'UAE', '12345', 'open'],
        ['Emma', 'Brown', 'emma.b@example.com', '+4455667788', 
         '2025-06-18', 'Air France', 'AF101', '101 River Road', 
         'Paris', 'France', '75001', 'open'],
        ['David', 'Jones', 'david.jones@example.com', '+9988776655', 
         '2025-06-20', 'Singapore Airlines', 'SQ202', '202 Mountain View', 
         'Singapore', 'Singapore', '123456', 'open']
    ]
    
    for row in sample_data:
        csv_writer.writerow(row)
    
    # Prepare the response
    output = csv_data.getvalue()
    csv_data.close()
    
    response = app.response_class(
        output,
        mimetype='text/csv',
        headers={"Content-Disposition": "attachment;filename=sample_flights.csv"}
    )
    
    return response

@app.route('/download_data')
def download_data():
    import csv, io
    from flask import Response

    # Get filters from query params
    date = request.args.get('date')
    admin = request.args.get('admin')
    status = request.args.get('status')
    
    # Start the query for Flight model
    query = Flight.query

    # Apply filters dynamically if they are provided
    if date:
        query = query.filter(Flight.flight_date == date)  # Exact date match
    
    if admin:
        query = query.join(User, Flight.user_id == User.id).filter(User.assigned_admin_id == admin)
    
    if status:
        query = query.filter(Flight.status == status)  # Exact status match

    # Fetch the filtered results
    flights = query.all()

    # Create an in-memory file to store CSV data
    output = io.StringIO()
    writer = csv.writer(output)

    # Write the header (defining column names from the Flight model)
    writer.writerow([
        'Flight ID', 'User ID', 'Leaving From', 'Destination', 'Connecting Flights',
        'Connecting Leaving From', 'Compensation Flight', 'Date', 'Disruption', 'Delay','Reason','Airline',
        'Airline 2','Airline 3', 'Flight Number', 'Flight Date', 'Address Line 1', 'Address Line 2', 
        'First Name', 'Last Name', 'First Name Passenger 2', 'Last Name Passenger 2', 
        'First Name Passenger 3', 'Last Name Passenger 3', 'First Name Passenger 4', 
        'Last Name Passenger 4', 'First Name Passenger 5', 'Last Name Passenger 5', 
        'Email', 'Group Travel', 'City', 'Postal Code', 'Country', 
        'Phone Number', 'Notification', 'Booking Number', 'Signature',
        'Contacted Airline', 'Incident Description', 'Status'
    ])

    # Write data rows for each flight object
    for flight in flights:
        writer.writerow([
            flight.id, flight.user_id, flight.leaving_from, flight.destination, flight.connecting_flights,
            flight.connecting_leaving_from, flight.compensation_flight, flight.date, flight.disruption, flight.delay,flight.reason, 
            flight.airline, flight.airline2,flight.airline3, flight.flight_number, flight.flight_date, flight.address_line1, 
            flight.address_line2, flight.first_name, flight.last_name, flight.first_name_passenger_2, 
            flight.last_name_passenger_2, flight.first_name_passenger_3, flight.last_name_passenger_3,
            flight.first_name_passenger_4, flight.last_name_passenger_4, flight.first_name_passenger_5,
            flight.last_name_passenger_5, flight.email, flight.group_travel, flight.city, flight.postal_code, 
            flight.country, flight.phone_number, flight.notification, flight.booking_number, 
            'Yes',flight.contacted_airline,
            flight.incident_description, flight.status
        ])

    # Set the pointer to the beginning of the stream
    output.seek(0)

    # Create the response with the correct MIME type for CSV
    return Response(output, mimetype="text/csv", 
                    headers={"Content-Disposition": "attachment;filename=flights.csv"})
#test

@app.route('/base')
def route():
    return render_template('base.html')


if __name__ == '__main__':
    init_db()  # Create database tables
    app.run(debug=True,host='0.0.0.0')












